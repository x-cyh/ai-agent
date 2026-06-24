from __future__ import annotations

import hashlib
import json
import re
import sys
import uuid
from datetime import datetime
from pathlib import Path

import streamlit as st

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SHELL_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from studio_shell.page_shell import page_shell
from studio_shell.shell_ui import (
    format_extra_context,
    inject_style,
    load_page_data,
    save_page_data,
    shared_data_path,
)

PAGE_NAME = "人生故事書"
OUTPUT_DIR = Path("story_outputs")
STORIES_DIR_NAME = "life_stories"

st.set_page_config(page_title="人生故事書", page_icon="📖", layout="wide")
inject_style()


PARENT_STORY_HINT = """
### 採訪流程提示

用「採訪式引導」幫您把一段人生故事整理成可以留給孩子的文字。

**左欄三個區塊**

- **故事主線**：一句話寫下最想留下的那段故事。
- **訪談進度筆記**：AI 每輪會自動 append 進度；您也可以手動補充關鍵字或場景。
- **已蒐集事實**：每行一筆 `key：value`（例如 `主角：媽媽`），AI 會自動 merge；您也可以手動編輯。

**右欄採訪流程**

1. 先確認主線：最想留下的那段故事。
2. 沿主線依序追問：背景 → 開端 → 衝突 → 決定 → 結果 → 意義 → 留給孩子的話。
3. 回答太短時，AI 會溫和補問細節；太抽象時，會拉回具體事件。
4. 資訊足夠後，跟 Agent 說「可以整理成故事了」，會產出 `.md` 檔到 `story_outputs/`。

**AI 自動紀錄格式**

Agent 每輪回應結尾會輸出兩個區塊，系統會自動寫進左欄：

- `【NOTES】...【/NOTES】` → append 到「訪談進度筆記」
- `【FACTS】...【/FACTS】` → merge 到「已蒐集事實」

**按鈕與功能**

- **選擇故事（下拉選單）**：切換到列表中的其他故事；切換後左欄三個區塊會自動載入該故事的內容。
- **➕ 新故事**：建立一個全新的空白故事，並自動切換為目前編輯對象。
- **💾 存為備份**：把目前故事「凍結」複製一份（標題加「（備份）」、狀態設為已完成），原故事不動，可繼續編輯。
- **✅ 標記完成 / ↩️ 標記進行中**：切換目前故事的狀態（採訪中 ↔ 已完成），內容不變。
- **🗑️ 刪除**：刪除目前故事（連同狀態檔）；刪除後會自動切到第一個剩下的故事。
- **故事標題**：可隨時改名；改完會自動同步到索引檔。
- **故事主線**：一句話寫下最想留下的那段故事。
- **訪談進度筆記**：AI 每輪會自動 append 進度；您也可以手動補充關鍵字或場景。
- **已蒐集事實**：每行一筆 `key：value`（例如 `主角：媽媽`），AI 會自動 merge；您也可以手動編輯。
- **訪談時間軸**：顯示最近 10 輪的對話摘要（時間、使用者訊息、筆記、事實）。支援三種來源：
  - 🤖 **Agent 自動記錄**：每輪 Agent 回應後自動 append。
  - ✍️ **手動編輯偵測**：您直接改「進度筆記」或「事實」時，系統會自動 append 一筆異動紀錄。
  - ➕ **手動新增**：在時間軸下方填寫訊息內容，按「➕ 加入時間軸」即可新增一筆。
  - 每筆右側有「🗑️ 刪除此筆」可單獨移除。
- **📄 匯出 Word / 📕 匯出 PDF**：把目前故事匯出成 `.docx` 或 `.pdf`，存到 `story_outputs/`。
- **⬇️ 下載 Word / ⬇️ 下載 PDF**：把剛匯出的檔案下載到本機。
- **給 Agent 的摘要**：把目前故事的狀態打包成一段文字，貼給 Agent 讓它接續採訪。

**狀態檔位置**

- 索引檔：`studio_shell/data/人生故事書.json`（故事列表 + 目前選擇）
- 每個故事：`studio_shell/data/life_stories/story_<id>.json`
""".strip()


# Agent 回應裡的結構化標記區塊（用正則解析）
NOTES_RE = re.compile(r"【NOTES】(.*?)【/NOTES】", re.DOTALL)
FACTS_RE = re.compile(r"【FACTS】(.*?)【/FACTS】", re.DOTALL)


# ─────────────────────────────────────────────────────────────
# 匯出：Word / PDF
# ─────────────────────────────────────────────────────────────

def _safe_filename(title: str) -> str:
    """把故事標題轉成檔名安全字串。"""
    name = re.sub(r"[\\/:*?\"<>|]", "_", title or "未命名故事")
    name = name.strip().strip(".")
    return name or "未命名故事"


# 七段式敘事骨架（背景 → 開端 → 衝突 → 決定 → 結果 → 意義 → 留給孩子的話）
STORY_SKELETON = [
    "背景",
    "開端",
    "衝突",
    "決定",
    "結果",
    "意義",
    "留給孩子的話",
]


def _compose_story(story: dict) -> str:
    """把事實 + 進度筆記編織成連貫敘事 Markdown 草稿。

    規則：
    - 標題用 story.title
    - 主線放最前面作為引子
    - 依「背景 → 開端 → 衝突 → 決定 → 結果 → 意義 → 留給孩子的話」順序，
      把對應的事實串成連貫敘事段落（不寫出小標題）
    - 沒對應事實的段落就跳過，不留「待補」提示
    - 沒被骨架對應到的事實，附加在主體之後
    - 結尾附上「事實索引」與「進度筆記」附錄
    """
    title = story.get("title") or "未命名故事"
    seed = (story.get("story_seed") or "").strip()
    facts = story.get("collected_facts") or {}
    notes = (story.get("progress_notes") or "").strip()

    # 把 facts 轉成小寫 key 對照表，方便不分大小寫配對
    facts_lower = {str(k).strip().lower(): (k, v) for k, v in facts.items() if k}

    lines: list[str] = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"*建立：{story.get('created_at', '')}　·　更新：{story.get('updated_at', '')}*")
    lines.append("")

    if seed:
        lines.append(seed)
        lines.append("")

    # 依七段式順序串成連貫敘事（不寫小標題）
    used_keys: set[str] = set()
    for heading in STORY_SKELETON:
        match_value = ""
        match_key = ""
        for low_key, (orig_key, value) in facts_lower.items():
            if low_key == heading.lower() or heading in str(orig_key):
                match_value = str(value).strip()
                match_key = str(orig_key)
                break
        if match_value:
            lines.append(match_value)
            lines.append("")
            used_keys.add(match_key)

    # 沒被骨架對應到的事實，附加在主體之後
    extra_facts = {k: v for k, v in facts.items() if k not in used_keys}
    if extra_facts:
        lines.append("---")
        lines.append("")
        for k, v in extra_facts.items():
            lines.append(f"**{k}**：{v}")
            lines.append("")

    # 附錄：事實索引
    if facts:
        lines.append("---")
        lines.append("")
        lines.append("## 附錄：事實索引")
        lines.append("")
        for k, v in facts.items():
            lines.append(f"- **{k}**：{v}")
        lines.append("")

    # 附錄：進度筆記
    if notes:
        lines.append("---")
        lines.append("")
        lines.append("## 附錄：訪談進度筆記")
        lines.append("")
        lines.append(notes)
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def _build_story_sections(story: dict) -> list[tuple[str, str]]:
    """把故事切成 (heading, body) 段落，給 docx/pdf 共用。"""
    sections: list[tuple[str, str]] = []

    title = story.get("title") or "未命名故事"
    sections.append(("標題", title))

    seed = story.get("story_seed", "")
    if seed:
        sections.append(("故事主線", seed))

    facts = story.get("collected_facts", {})
    if isinstance(facts, dict) and facts:
        lines = [f"{k}：{v}" for k, v in facts.items()]
        sections.append(("已蒐集事實", "\n".join(lines)))

    notes = story.get("progress_notes", "")
    if notes:
        sections.append(("訪談進度筆記", notes))

    timeline = story.get("timeline", [])
    if isinstance(timeline, list) and timeline:
        lines = []
        for entry in timeline:
            ts = entry.get("ts", "")
            user = entry.get("user", "")
            notes_added = entry.get("notes_added", "")
            facts_added = entry.get("facts_added", [])
            line = f"[{ts}] 使用者：{user}"
            if notes_added:
                line += f"\n　→ 筆記：{notes_added}"
            if facts_added:
                line += f"\n　→ 事實：{', '.join(facts_added)}"
            lines.append(line)
        sections.append(("訪談時間軸", "\n\n".join(lines)))

    return sections


def _export_to_docx(story: dict, output_path: Path) -> None:
    """把故事匯出成 .docx。"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 標題
    title_para = doc.add_heading(story.get("title") or "未命名故事", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 中繼資訊
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"建立：{story.get('created_at', '')}　·　"
        f"更新：{story.get('updated_at', '')}　·　"
        f"狀態：{'已完成' if story.get('status') == 'completed' else '採訪中'}"
    )
    meta_run.font.size = Pt(9)

    doc.add_paragraph()  # 空行

    for heading, body in _build_story_sections(story):
        if heading == "標題":
            continue  # 已加過
        doc.add_heading(heading, level=1)
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _export_composed_to_docx(markdown_text: str, title: str, output_path: Path) -> None:
    """把「完整故事」Markdown 草稿匯出成 .docx（簡單逐行轉換）。"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title_para = doc.add_heading(title or "未命名故事", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            # 已加過標題，跳過
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("---"):
            # 分隔線：用一個空段
            doc.add_paragraph("─" * 30)
            continue
        if line.startswith("*") and line.endswith("*") and len(line) > 2:
            # 斜體中繼資訊
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*").strip())
            run.italic = True
            run.font.size = Pt(9)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _export_composed_to_pdf(markdown_text: str, title: str, output_path: Path) -> None:
    """把「完整故事」Markdown 草稿匯出成 .pdf（支援中文）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title or "未命名故事",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=22,
        leading=28,
        alignment=1,
        spaceAfter=12,
    )
    h1_style = ParagraphStyle(
        "H1CN",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=22,
        spaceBefore=14,
        spaceAfter=8,
    )
    h2_style = ParagraphStyle(
        "H2CN",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=13,
        leading=18,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "MetaCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9,
        leading=14,
        alignment=1,
        textColor="#666666",
        spaceAfter=12,
    )

    flow = []
    flow.append(Paragraph(_escape_xml(title or "未命名故事"), title_style))

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            flow.append(Spacer(1, 0.3 * cm))
            continue
        if line.startswith("# "):
            continue  # 已加過標題
        if line.startswith("## "):
            flow.append(Paragraph(_escape_xml(line[3:].strip()), h1_style))
            continue
        if line.startswith("### "):
            flow.append(Paragraph(_escape_xml(line[4:].strip()), h2_style))
            continue
        if line.startswith("---"):
            flow.append(Spacer(1, 0.4 * cm))
            continue
        if line.startswith("*") and line.endswith("*") and len(line) > 2:
            flow.append(Paragraph(_escape_xml(line.strip("*").strip()), meta_style))
            continue
        if line.startswith("- "):
            flow.append(Paragraph("• " + _escape_xml(line[2:].strip()), body_style))
            continue
        flow.append(Paragraph(_escape_xml(line), body_style))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(flow)


def _export_to_pdf(story: dict, output_path: Path) -> None:
    """把故事匯出成 .pdf（支援中文：使用 reportlab + CID 字型）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # 註冊繁體中文 CID 字型（reportlab 內建，不需要外部 ttf）
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=story.get("title") or "未命名故事",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=22,
        leading=28,
        alignment=1,  # CENTER
        spaceAfter=12,
    )
    h1_style = ParagraphStyle(
        "H1CN",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=22,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "MetaCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9,
        leading=14,
        alignment=1,
        textColor="#666666",
        spaceAfter=12,
    )

    story_flow = []

    title = story.get("title") or "未命名故事"
    story_flow.append(Paragraph(_escape_xml(title), title_style))

    meta_text = (
        f"建立：{story.get('created_at', '')}　·　"
        f"更新：{story.get('updated_at', '')}　·　"
        f"狀態：{'已完成' if story.get('status') == 'completed' else '採訪中'}"
    )
    story_flow.append(Paragraph(_escape_xml(meta_text), meta_style))
    story_flow.append(Spacer(1, 0.5 * cm))

    for heading, body in _build_story_sections(story):
        if heading == "標題":
            continue
        story_flow.append(Paragraph(_escape_xml(heading), h1_style))
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            # 把換行轉成 <br/>
            line_xml = _escape_xml(line).replace("\n", "<br/>")
            story_flow.append(Paragraph(line_xml, body_style))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story_flow)


def _escape_xml(text: str) -> str:
    """Escape XML special chars for reportlab Paragraph."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


# ─────────────────────────────────────────────────────────────
# 資料層：故事索引 + 每個故事一份檔
# ─────────────────────────────────────────────────────────────

def _stories_dir(shell_root: Path) -> Path:
    return shell_root / "data" / STORIES_DIR_NAME


def _new_story_id() -> str:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    short = uuid.uuid4().hex[:6]
    return f"story_{stamp}_{short}"


def _story_path(story_id: str, shell_root: Path) -> Path:
    return _stories_dir(shell_root) / f"{story_id}.json"


def _load_index(shell_root: Path) -> dict:
    """讀取索引檔，自動 migrate 舊格式。"""
    index_path = shared_data_path(PAGE_NAME, shell_root=shell_root)
    if not index_path.is_file():
        return {"current_story_id": "", "stories": []}

    try:
        raw = json.loads(index_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"current_story_id": "", "stories": []}

    if not isinstance(raw, dict):
        return {"current_story_id": "", "stories": []}

    # 已是新格式
    if "stories" in raw and isinstance(raw["stories"], list):
        return raw

    # 舊格式 migrate：把現有資料包成一個 story
    legacy_id = f"story_legacy_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    legacy_story = {
        "id": legacy_id,
        "title": str(raw.get("story_seed", "") or "未命名故事")[:50] or "未命名故事",
        "story_seed": str(raw.get("story_seed", "") or ""),
        "progress_notes": str(raw.get("progress_notes", "") or ""),
        "collected_facts": raw.get("collected_facts", {}) if isinstance(raw.get("collected_facts"), dict) else {},
        "timeline": raw.get("timeline", []) if isinstance(raw.get("timeline"), list) else [],
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "in_progress",
        "migrated_from_legacy": True,
    }
    _stories_dir(shell_root).mkdir(parents=True, exist_ok=True)
    _story_path(legacy_id, shell_root).write_text(
        json.dumps(legacy_story, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    new_index = {
        "current_story_id": legacy_id,
        "stories": [{
            "id": legacy_id,
            "title": legacy_story["title"],
            "created_at": legacy_story["created_at"],
            "updated_at": legacy_story["updated_at"],
            "status": "in_progress",
        }],
    }
    index_path.write_text(
        json.dumps(new_index, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return new_index


def _save_index(index: dict, shell_root: Path) -> None:
    index_path = shared_data_path(PAGE_NAME, shell_root=shell_root)
    index_path.parent.mkdir(parents=True, exist_ok=True)
    index_path.write_text(
        json.dumps(index, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _load_story(story_id: str, shell_root: Path) -> dict | None:
    path = _story_path(story_id, shell_root)
    if not path.is_file():
        return None
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    return raw if isinstance(raw, dict) else None


def _save_story(story: dict, shell_root: Path) -> None:
    story["updated_at"] = datetime.now().isoformat(timespec="seconds")
    _stories_dir(shell_root).mkdir(parents=True, exist_ok=True)
    _story_path(story["id"], shell_root).write_text(
        json.dumps(story, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _create_story(title: str, shell_root: Path) -> dict:
    now = datetime.now().isoformat(timespec="seconds")
    story = {
        "id": _new_story_id(),
        "title": title or "未命名故事",
        "story_seed": "",
        "progress_notes": "",
        "collected_facts": {},
        "timeline": [],
        "created_at": now,
        "updated_at": now,
        "status": "in_progress",
    }
    _save_story(story, shell_root)
    return story


def _delete_story(story_id: str, shell_root: Path) -> None:
    path = _story_path(story_id, shell_root)
    if path.is_file():
        path.unlink()


# ─────────────────────────────────────────────────────────────
# 解析 Agent 回應
# ─────────────────────────────────────────────────────────────

def _parse_facts_block(block: str) -> dict[str, str]:
    facts: dict[str, str] = {}
    for line in block.strip().splitlines():
        line = line.strip()
        if not line or "：" not in line:
            continue
        key, value = line.split("：", 1)
        key = key.strip()
        value = value.strip()
        if key:
            facts[key] = value
    return facts


def _on_agent_reply(answer: str, user_text: str) -> None:
    """Agent 回應後自動更新「目前選擇的故事」。"""
    sig = hashlib.md5(answer.encode("utf-8")).hexdigest()
    if st.session_state.get("_life_story_last_sig") == sig:
        return
    st.session_state["_life_story_last_sig"] = sig

    notes_match = NOTES_RE.search(answer)
    facts_match = FACTS_RE.search(answer)
    new_notes = notes_match.group(1).strip() if notes_match else ""
    new_facts = _parse_facts_block(facts_match.group(1)) if facts_match else {}

    if not new_notes and not new_facts:
        return

    index = _load_index(SHELL_ROOT)
    story_id = index.get("current_story_id", "")
    if not story_id:
        return

    story = _load_story(story_id, SHELL_ROOT)
    if story is None:
        return

    existing_notes = str(story.get("progress_notes", "") or "")
    if new_notes:
        merged_notes = (existing_notes + "\n\n" + new_notes).strip() if existing_notes else new_notes
    else:
        merged_notes = existing_notes

    existing_facts = story.get("collected_facts", {})
    if not isinstance(existing_facts, dict):
        existing_facts = {}
    merged_facts = {**existing_facts, **new_facts}

    timeline = story.get("timeline", [])
    if not isinstance(timeline, list):
        timeline = []
    timeline.append({
        "ts": datetime.now().isoformat(timespec="seconds"),
        "user": user_text[:200],
        "agent_excerpt": answer[:200],
        "notes_added": new_notes[:200],
        "facts_added": list(new_facts.keys()),
    })

    story["progress_notes"] = merged_notes
    story["collected_facts"] = merged_facts
    story["timeline"] = timeline
    _save_story(story, SHELL_ROOT)

    # 同步更新索引裡的 updated_at
    for entry in index.get("stories", []):
        if entry.get("id") == story_id:
            entry["updated_at"] = story["updated_at"]
            break
    _save_index(index, SHELL_ROOT)


# ─────────────────────────────────────────────────────────────
# UI helpers
# ─────────────────────────────────────────────────────────────

def _facts_to_text(facts: dict[str, str]) -> str:
    return "\n".join(f"{k}：{v}" for k, v in facts.items())


def _text_to_facts(text: str) -> dict[str, str]:
    facts: dict[str, str] = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or "：" not in line:
            continue
        key, value = line.split("：", 1)
        key = key.strip()
        value = value.strip()
        if key:
            facts[key] = value
    return facts


def _story_label(entry: dict) -> str:
    title = entry.get("title", "未命名")
    status = entry.get("status", "in_progress")
    updated = entry.get("updated_at", "")[:16].replace("T", " ")
    status_icon = "✅" if status == "completed" else "📝"
    return f"{status_icon} {title}（{updated}）"


# ─────────────────────────────────────────────────────────────
# 主畫面
# ─────────────────────────────────────────────────────────────

def render_main() -> str:
    index = _load_index(SHELL_ROOT)
    stories = index.get("stories", [])
    current_id = index.get("current_story_id", "")

    # 確保 current_id 有效
    if current_id and not _load_story(current_id, SHELL_ROOT):
        current_id = ""
        index["current_story_id"] = ""
    if not current_id and stories:
        current_id = stories[0]["id"]
        index["current_story_id"] = current_id
        _save_index(index, SHELL_ROOT)

    # ── 故事管理區塊 ──
    st.markdown("#### 📚 故事管理")

    # 「建立新故事」表單狀態（按下 ➕ 後展開）
    new_form_key = "_life_story_new_form_open"

    def _open_new_form() -> None:
        st.session_state[new_form_key] = True
        st.session_state["_life_story_new_title_input"] = ""

    def _close_new_form() -> None:
        st.session_state[new_form_key] = False
        st.session_state["_life_story_new_title_input"] = ""

    def _create_and_switch(title: str) -> None:
        story = _create_story(title, SHELL_ROOT)
        index["stories"].append({
            "id": story["id"],
            "title": story["title"],
            "created_at": story["created_at"],
            "updated_at": story["updated_at"],
            "status": "in_progress",
        })
        index["current_story_id"] = story["id"]
        _save_index(index, SHELL_ROOT)
        _close_new_form()
        st.rerun()

    if not stories:
        st.info("還沒有故事。點下方「➕ 新故事」開始第一段。")
        if not st.session_state.get(new_form_key, False):
            col_new, _ = st.columns([1, 5])
            if col_new.button("➕ 新故事", type="primary", use_container_width=True):
                _open_new_form()
                st.rerun()
        else:
            # 展開輸入表單
            with st.container():
                st.markdown("##### ✏️ 為新故事命名")
                new_title_input = st.text_input(
                    "故事名稱",
                    placeholder="例如：媽媽北上工作",
                    key="_life_story_new_title_input",
                )
                col_ok, col_cancel = st.columns([1, 1])
                ok_clicked = col_ok.button(
                    "✅ 建立",
                    type="primary",
                    use_container_width=True,
                    key="_life_story_new_ok",
                )
                cancel_clicked = col_cancel.button(
                    "❌ 取消",
                    use_container_width=True,
                    key="_life_story_new_cancel",
                )
                if ok_clicked:
                    title = (new_title_input or "").strip()
                    if not title:
                        st.warning("請輸入故事名稱")
                    else:
                        _create_and_switch(title)
                if cancel_clicked:
                    _close_new_form()
                    st.rerun()
        return ""

    # 選擇器
    story_labels = {entry["id"]: _story_label(entry) for entry in stories}
    selected_id = st.selectbox(
        "選擇故事",
        options=list(story_labels.keys()),
        index=list(story_labels.keys()).index(current_id) if current_id in story_labels else 0,
        format_func=lambda sid: story_labels.get(sid, sid),
        key="life_story_picker",
    )

    # 切換故事
    if selected_id != current_id:
        index["current_story_id"] = selected_id
        _save_index(index, SHELL_ROOT)
        st.rerun()

    # 操作按鈕
    current_story = _load_story(current_id, SHELL_ROOT) or {}
    is_completed = current_story.get("status") == "completed"

    btn_col1, btn_col2, btn_col3, btn_col4 = st.columns(4)
    new_clicked = btn_col1.button("➕ 新故事", use_container_width=True)
    snapshot_clicked = btn_col2.button("💾 存為備份", use_container_width=True)
    complete_label = "↩️ 標記進行中" if is_completed else "✅ 標記完成"
    complete_clicked = btn_col3.button(complete_label, use_container_width=True)
    delete_clicked = btn_col4.button("🗑️ 刪除", type="secondary", use_container_width=True)

    if new_clicked:
        _open_new_form()
        st.rerun()

    # 「建立新故事」表單（按下 ➕ 後展開）
    if st.session_state.get(new_form_key, False):
        with st.container():
            st.markdown("##### ✏️ 為新故事命名")
            new_title_input = st.text_input(
                "故事名稱",
                placeholder="例如：媽媽北上工作",
                key="_life_story_new_title_input",
            )
            col_ok, col_cancel = st.columns([1, 1])
            ok_clicked = col_ok.button(
                "✅ 建立",
                type="primary",
                use_container_width=True,
                key="_life_story_new_ok",
            )
            cancel_clicked = col_cancel.button(
                "❌ 取消",
                use_container_width=True,
                key="_life_story_new_cancel",
            )
            if ok_clicked:
                title = (new_title_input or "").strip()
                if not title:
                    st.warning("請輸入故事名稱")
                else:
                    _create_and_switch(title)
            if cancel_clicked:
                _close_new_form()
                st.rerun()

    if snapshot_clicked and current_id:
        story = _load_story(current_id, SHELL_ROOT)
        if story:
            snapshot_id = f"{current_id}_snapshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            snapshot = dict(story)
            snapshot["id"] = snapshot_id
            snapshot["title"] = f"{story.get('title', '')}（備份）"
            snapshot["status"] = "completed"
            _save_story(snapshot, SHELL_ROOT)
            index["stories"].append({
                "id": snapshot_id,
                "title": snapshot["title"],
                "created_at": snapshot["created_at"],
                "updated_at": snapshot["updated_at"],
                "status": "completed",
            })
            _save_index(index, SHELL_ROOT)
            st.success(f"已存備份：{snapshot['title']}")

    if complete_clicked and current_id:
        story = _load_story(current_id, SHELL_ROOT)
        if story:
            new_status = "in_progress" if story.get("status") == "completed" else "completed"
            story["status"] = new_status
            _save_story(story, SHELL_ROOT)
            for entry in index["stories"]:
                if entry["id"] == current_id:
                    entry["status"] = new_status
                    break
            _save_index(index, SHELL_ROOT)
            st.rerun()

    if delete_clicked and current_id:
        _delete_story(current_id, SHELL_ROOT)
        index["stories"] = [e for e in index["stories"] if e["id"] != current_id]
        if index["stories"]:
            index["current_story_id"] = index["stories"][0]["id"]
        else:
            index["current_story_id"] = ""
        _save_index(index, SHELL_ROOT)
        st.rerun()

    # ── 載入目前故事 ──
    story = current_story

    # ── 採訪流程提示（折疊） ──
    with st.expander("📖 採訪流程與使用說明", expanded=False):
        st.markdown(PARENT_STORY_HINT)

    # ── 故事標題（可改） ──
    new_title = st.text_input(
        "故事標題",
        value=story.get("title", ""),
        placeholder="例如：媽媽北上工作",
        key="life_story_title",
    )
    if new_title != story.get("title", ""):
        story["title"] = new_title
        for entry in index["stories"]:
            if entry["id"] == current_id:
                entry["title"] = new_title
                break
        _save_story(story, SHELL_ROOT)
        _save_index(index, SHELL_ROOT)

    # ── 故事主線 ──
    st.markdown("#### 故事主線")
    story_seed = st.text_input(
        "如果今天只能先留下一段故事，您最想從哪一件事開始說起？",
        value=story.get("story_seed", ""),
        placeholder="例如：我媽年輕時一個人北上工作的那段日子",
        key="life_story_seed",
    )

    # ── 訪談進度筆記 ──
    st.markdown("#### 訪談進度筆記")
    st.caption("AI 會自動 append 進度；您也可以直接編輯補充。")
    progress_notes = st.text_area(
        "目前聊到什麼階段？有什麼想先記下來的關鍵字或場景？",
        value=story.get("progress_notes", ""),
        placeholder="例如：已確認主線、剛問完背景、接下來要問最大衝突",
        height=160,
        key="life_story_notes",
    )

    # ── 已蒐集事實 ──
    st.markdown("#### 已蒐集事實")
    st.caption("每行一筆「key：value」，AI 會自動 merge；您也可以手動編輯。")
    existing_facts = story.get("collected_facts", {})
    if not isinstance(existing_facts, dict):
        existing_facts = {}
    facts_text = st.text_area(
        "事實清單",
        value=_facts_to_text(existing_facts),
        placeholder="例如：\n主角：媽媽\n年代：1970 年代\n地點：台北",
        height=160,
        key="life_story_facts_editor",
    )
    edited_facts = _text_to_facts(facts_text)

    # ── 訪談時間軸 ──
    timeline = story.get("timeline", [])
    if not isinstance(timeline, list):
        timeline = []

    # ── 自動偵測「手動編輯」並 append 時間軸 ──
    # 用 session_state 記住上次存檔的快照，比對差異
    notes_sig_key = f"_life_story_last_notes_sig_{current_id}"
    facts_sig_key = f"_life_story_last_facts_sig_{current_id}"
    notes_sig_now = hashlib.md5(progress_notes.encode("utf-8")).hexdigest()
    facts_sig_now = hashlib.md5(facts_text.encode("utf-8")).hexdigest()

    notes_prev = st.session_state.get(notes_sig_key)
    facts_prev = st.session_state.get(facts_sig_key)

    # 第一次進入此故事：只記錄快照，不 append
    if notes_prev is None:
        st.session_state[notes_sig_key] = notes_sig_now
    elif notes_prev != notes_sig_now:
        # 進度筆記被改過了
        timeline.append({
            "ts": datetime.now().isoformat(timespec="seconds"),
            "user": "（手動編輯進度筆記）",
            "agent_excerpt": "",
            "notes_added": f"更新進度筆記（{len(progress_notes)} 字）",
            "facts_added": [],
            "source": "manual_edit",
        })
        st.session_state[notes_sig_key] = notes_sig_now

    if facts_prev is None:
        st.session_state[facts_sig_key] = facts_sig_now
    elif facts_prev != facts_sig_now:
        # 事實被改過了：算出新增/修改的 keys
        old_facts = _text_to_facts(st.session_state.get(f"_life_story_last_facts_text_{current_id}", ""))
        new_facts = edited_facts
        added_keys = [k for k in new_facts.keys() if k not in old_facts]
        changed_keys = [k for k in new_facts.keys() if k in old_facts and old_facts[k] != new_facts[k]]
        removed_keys = [k for k in old_facts.keys() if k not in new_facts]
        diff_keys = added_keys + changed_keys + removed_keys
        if diff_keys:
            timeline.append({
                "ts": datetime.now().isoformat(timespec="seconds"),
                "user": "（手動編輯事實）",
                "agent_excerpt": "",
                "notes_added": f"異動：{', '.join(diff_keys)}",
                "facts_added": diff_keys,
                "source": "manual_edit",
            })
        st.session_state[facts_sig_key] = facts_sig_now
        st.session_state[f"_life_story_last_facts_text_{current_id}"] = facts_text

    # 初始化事實快照（給下一輪比對用）
    if f"_life_story_last_facts_text_{current_id}" not in st.session_state:
        st.session_state[f"_life_story_last_facts_text_{current_id}"] = facts_text

    with st.expander(f"訪談時間軸（{len(timeline)} 輪，顯示最近 10 輪）", expanded=False):
        if not timeline:
            st.caption("（尚無紀錄）")
        else:
            # 顯示最近 10 輪（由新到舊）
            recent = timeline[-10:][::-1]
            for idx, entry in enumerate(recent):
                ts = entry.get("ts", "")
                user = entry.get("user", "")
                notes = entry.get("notes_added", "")
                facts_added = entry.get("facts_added", [])
                source = entry.get("source", "agent")
                source_icon = "✍️" if source == "manual_edit" else ("🤖" if source == "agent" else "📝")
                if source == "manual_add":
                    source_icon = "➕"
                st.markdown(f"{source_icon} **[{ts}]** {user}")
                if notes:
                    st.markdown(f"　→ 筆記：{notes}")
                if facts_added:
                    st.markdown(f"　→ 事實：{', '.join(facts_added)}")
                # 刪除單筆按鈕
                real_idx = len(timeline) - 1 - idx
                del_key = f"del_timeline_{current_id}_{real_idx}"
                if st.button("🗑️ 刪除此筆", key=del_key, use_container_width=False):
                    timeline.pop(real_idx)
                    story["timeline"] = timeline
                    _save_story(story, SHELL_ROOT)
                    # 重置快照，避免下次又把同一份編輯當成新異動
                    st.session_state[notes_sig_key] = hashlib.md5(progress_notes.encode("utf-8")).hexdigest()
                    st.session_state[facts_sig_key] = hashlib.md5(facts_text.encode("utf-8")).hexdigest()
                    st.rerun()

        st.divider()
        # ── 手動新增記錄 ──
        st.markdown("**➕ 手動新增一筆紀錄**")
        manual_col1, manual_col2 = st.columns([3, 1])
        manual_msg = manual_col1.text_input(
            "訊息內容",
            placeholder="例如：訪談結束，今天聊到決定北上那一刻",
            key=f"manual_timeline_msg_{current_id}",
        )
        manual_notes = manual_col2.text_input(
            "筆記（可選）",
            placeholder="關鍵字",
            key=f"manual_timeline_notes_{current_id}",
        )
        if st.button("➕ 加入時間軸", key=f"btn_manual_timeline_{current_id}", use_container_width=True):
            if manual_msg.strip() or manual_notes.strip():
                timeline.append({
                    "ts": datetime.now().isoformat(timespec="seconds"),
                    "user": manual_msg.strip() or "（手動記錄）",
                    "agent_excerpt": "",
                    "notes_added": manual_notes.strip(),
                    "facts_added": [],
                    "source": "manual_add",
                })
                story["timeline"] = timeline
                _save_story(story, SHELL_ROOT)
                st.success("已加入時間軸")
                st.rerun()
            else:
                st.warning("請至少填寫「訊息內容」或「筆記」其中一欄")

    # ── 存檔 ──
    story["story_seed"] = story_seed
    story["progress_notes"] = progress_notes
    story["collected_facts"] = edited_facts
    story["timeline"] = timeline
    _save_story(story, SHELL_ROOT)

    # 同步索引 updated_at
    for entry in index["stories"]:
        if entry["id"] == current_id:
            entry["updated_at"] = story["updated_at"]
            break
    _save_index(index, SHELL_ROOT)

    # ── 匯出文件 ──
    st.markdown("#### 📤 匯出文件")
    st.caption("把目前故事匯出成 Word 或 PDF，存到 `story_outputs/` 並提供下載。")

    export_dir = PROJECT_ROOT / OUTPUT_DIR
    safe_title = _safe_filename(story.get("title", "未命名故事"))
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"{safe_title}_{stamp}.docx"
    pdf_filename = f"{safe_title}_{stamp}.pdf"
    docx_path = export_dir / docx_filename
    pdf_path = export_dir / pdf_filename

    exp_col1, exp_col2 = st.columns(2)

    with exp_col1:
        if st.button("📄 匯出 Word（.docx）", use_container_width=True):
            try:
                _export_to_docx(story, docx_path)
                st.session_state["_life_story_last_docx"] = str(docx_path)
                st.success(f"已匯出：{docx_filename}")
            except Exception as exc:
                st.error(f"Word 匯出失敗：`{exc}`")

    with exp_col2:
        if st.button("📕 匯出 PDF（.pdf）", use_container_width=True):
            try:
                _export_to_pdf(story, pdf_path)
                st.session_state["_life_story_last_pdf"] = str(pdf_path)
                st.success(f"已匯出：{pdf_filename}")
            except Exception as exc:
                st.error(f"PDF 匯出失敗：`{exc}`")

    # 下載按鈕（檔案存在才顯示）
    dl_col1, dl_col2 = st.columns(2)
    if docx_path.is_file():
        with dl_col1:
            st.download_button(
                "⬇️ 下載 Word",
                data=docx_path.read_bytes(),
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    if pdf_path.is_file():
        with dl_col2:
            st.download_button(
                "⬇️ 下載 PDF",
                data=pdf_path.read_bytes(),
                file_name=pdf_filename,
                mime="application/pdf",
                use_container_width=True,
            )

    st.divider()

    # ── 完整故事預覽（七段式編織） ──
    st.markdown("#### 📖 完整故事預覽")
    st.caption(
        "把「故事主線 + 已蒐集事實 + 訪談進度筆記」依「背景 → 開端 → 衝突 → 決定 → 結果 → 意義 → 留給孩子的話」"
        "的順序串成連貫敘事（不寫小標題）；可直接編輯後匯出 Word / PDF / Markdown。"
    )

    # 自動生成草稿（若使用者尚未編輯過，或按下「重新生成」）
    auto_draft = _compose_story(story)
    edited_key = "life_story_composed_md"
    regen_key = "life_story_composed_regen"

    if edited_key not in st.session_state:
        st.session_state[edited_key] = auto_draft

    regen_col, _ = st.columns([1, 5])
    if regen_col.button("🔄 重新從事實生成草稿", use_container_width=True):
        st.session_state[edited_key] = auto_draft
        st.session_state[regen_key] = datetime.now().isoformat(timespec="seconds")

    composed_md = st.text_area(
        "故事草稿（可直接編輯）",
        value=st.session_state[edited_key],
        height=320,
        key=edited_key,
    )

    # 預覽（用 Markdown 渲染）
    with st.expander("👀 預覽渲染結果", expanded=False):
        st.markdown(composed_md)

    # 匯出按鈕
    composed_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    composed_safe = _safe_filename(story.get("title", "未命名故事"))
    composed_docx_name = f"{composed_safe}_完整故事_{composed_stamp}.docx"
    composed_pdf_name = f"{composed_safe}_完整故事_{composed_stamp}.pdf"
    composed_md_name = f"{composed_safe}_完整故事_{composed_stamp}.md"
    composed_docx_path = export_dir / composed_docx_name
    composed_pdf_path = export_dir / composed_pdf_name
    composed_md_path = export_dir / composed_md_name

    exp_c1, exp_c2, exp_c3 = st.columns(3)
    with exp_c1:
        if st.button("📄 匯出 Word", key="btn_export_composed_docx", use_container_width=True):
            try:
                _export_composed_to_docx(composed_md, story.get("title", "未命名故事"), composed_docx_path)
                st.session_state["_life_story_last_composed_docx"] = str(composed_docx_path)
                st.success(f"已匯出：{composed_docx_name}")
            except Exception as exc:
                st.error(f"Word 匯出失敗：`{exc}`")
    with exp_c2:
        if st.button("📕 匯出 PDF", key="btn_export_composed_pdf", use_container_width=True):
            try:
                _export_composed_to_pdf(composed_md, story.get("title", "未命名故事"), composed_pdf_path)
                st.session_state["_life_story_last_composed_pdf"] = str(composed_pdf_path)
                st.success(f"已匯出：{composed_pdf_name}")
            except Exception as exc:
                st.error(f"PDF 匯出失敗：`{exc}`")
    with exp_c3:
        if st.button("📝 匯出 Markdown", key="btn_export_composed_md", use_container_width=True):
            try:
                composed_md_path.parent.mkdir(parents=True, exist_ok=True)
                composed_md_path.write_text(composed_md, encoding="utf-8")
                st.session_state["_life_story_last_composed_md"] = str(composed_md_path)
                st.success(f"已匯出：{composed_md_name}")
            except Exception as exc:
                st.error(f"Markdown 匯出失敗：`{exc}`")

    # 下載按鈕（檔案存在才顯示）
    dl_c1, dl_c2, dl_c3 = st.columns(3)
    if composed_docx_path.is_file():
        with dl_c1:
            st.download_button(
                "⬇️ 下載 Word",
                data=composed_docx_path.read_bytes(),
                file_name=composed_docx_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_composed_docx",
            )
    if composed_pdf_path.is_file():
        with dl_c2:
            st.download_button(
                "⬇️ 下載 PDF",
                data=composed_pdf_path.read_bytes(),
                file_name=composed_pdf_name,
                mime="application/pdf",
                use_container_width=True,
                key="dl_composed_pdf",
            )
    if composed_md_path.is_file():
        with dl_c3:
            st.download_button(
                "⬇️ 下載 Markdown",
                data=composed_md_path.read_bytes(),
                file_name=composed_md_name,
                mime="text/markdown",
                use_container_width=True,
                key="dl_composed_md",
            )

    st.divider()

    # ── 給 Agent 的摘要 ──
    st.markdown("#### 給 Agent 的摘要")
    extra = format_extra_context(
        PAGE_NAME,
        共享資料檔=str(shared_data_path(PAGE_NAME, shell_root=SHELL_ROOT)),
        目前故事=story.get("title", "未命名"),
        故事主線=story_seed or "（尚未設定）",
        訪談進度筆記=progress_notes or "（空白）",
        已蒐集事實=_facts_to_text(edited_facts) or "（尚無）",
        訪談輪數=str(len(timeline)),
        故事輸出目錄=str(OUTPUT_DIR),
        輸出檔名慣例="{timestamp}_{故事標題}.md",
        採訪流程="請嚴格遵循 parent-story skill：先確認主線 → 背景 → 開端 → 衝突 → 決定 → 結果 → 意義 → 留給孩子的話；每次只問一題，優先補充細節。",
        自動紀錄規則="每輪回應結尾請輸出一個【NOTES】區塊（簡述本輪進度與關鍵字）與一個【FACTS】區塊（每行 key：value，列出本輪新蒐集到的事實），系統會自動 append 到目前故事的狀態檔。資訊足夠後再整理輸出 story_outputs/ 的 .md 檔案。",
    )
    st.code(extra, language="text")

    return extra


page_shell(
    "人生故事書",
    "留下最想說給孩子聽的一段人生故事。",
    render_main,
    page_name=PAGE_NAME,
    on_assistant_reply=_on_agent_reply,
)